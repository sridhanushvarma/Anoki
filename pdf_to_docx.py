#!/usr/bin/env python3
"""
PDF to DOCX Converter Script
This script converts PDF files to DOCX format using pdf2docx library.
Includes comprehensive error handling and validation.
"""

import sys
import os
import logging
from pathlib import Path

# Try to import pdf2docx, fall back to alternative if needed
try:
    from pdf2docx.converter import Converter
    HAS_PDF2DOCX = True
except ImportError:
    HAS_PDF2DOCX = False

# Try to import PyMuPDF as fallback
try:
    import fitz
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

# Try to import python-docx for creating DOCX files
try:
    from docx import Document
    from docx.shared import Pt, Inches
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    stream=sys.stderr
)
logger = logging.getLogger(__name__)


def validate_pdf_file(input_path):
    """
    Validate that the input file is a valid PDF.

    Args:
        input_path (str): Path to the PDF file

    Returns:
        tuple: (is_valid, error_message)
    """
    try:
        # Check if file exists
        if not os.path.exists(input_path):
            return False, f"Input file '{input_path}' does not exist"

        # Check if it's a file (not a directory)
        if not os.path.isfile(input_path):
            return False, f"'{input_path}' is not a file"

        # Check file extension
        if not input_path.lower().endswith('.pdf'):
            return False, "Input file must be a PDF"

        # Check file size (max 100MB)
        file_size = os.path.getsize(input_path)
        max_size = 100 * 1024 * 1024  # 100MB
        if file_size > max_size:
            return False, f"File size ({file_size / 1024 / 1024:.2f}MB) exceeds maximum (100MB)"

        # Check if file is readable
        if not os.access(input_path, os.R_OK):
            return False, "Input file is not readable"

        # Check if it's a valid PDF by reading the header
        try:
            with open(input_path, 'rb') as f:
                header = f.read(4)
                if header != b'%PDF':
                    return False, "File does not appear to be a valid PDF (invalid header)"
        except Exception as e:
            return False, f"Error reading file: {str(e)}"

        return True, None

    except Exception as e:
        return False, f"Validation error: {str(e)}"


def convert_pdf_to_docx_fallback(input_path, output_path, page_from=None, page_to=None):
    """
    Fallback conversion using PyMuPDF and python-docx.
    Extracts text from PDF and creates a DOCX file.
    """
    if not HAS_PYMUPDF or not HAS_PYTHON_DOCX:
        logger.error("PyMuPDF and python-docx are required for fallback conversion")
        return False

    try:
        # Open PDF
        pdf_doc = fitz.open(input_path)
        num_pages = pdf_doc.page_count

        logger.info(f"PDF has {num_pages} pages")

        # Determine page range
        start_page = 0
        end_page = num_pages

        if page_from is not None:
            start_page = max(0, page_from - 1)
        if page_to is not None:
            end_page = min(page_to, num_pages)

        logger.info(f"Converting pages {start_page + 1} to {end_page}")

        # Create DOCX document
        doc = Document()

        # Extract text from each page
        for page_num in range(start_page, end_page):
            page = pdf_doc[page_num]
            text = page.get_text()

            if text.strip():
                # Add page break if not first page
                if page_num > start_page:
                    doc.add_page_break()

                # Add page header
                heading = doc.add_heading(f'Page {page_num + 1}', level=2)
                heading.style = 'Heading 2'

                # Add text content
                for line in text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)

        # Save DOCX
        doc.save(output_path)
        pdf_doc.close()

        logger.info(f"Fallback conversion completed successfully!")
        return True

    except Exception as e:
        logger.error(f"Fallback conversion error: {str(e)}", exc_info=True)
        return False


def convert_pdf_to_docx(input_path, output_path, maintain_formatting=True, extract_images=True, page_from=None, page_to=None):
    """
    Convert PDF to DOCX with comprehensive error handling.
    Uses pdf2docx if available, falls back to PyMuPDF + python-docx.

    Args:
        input_path (str): Path to input PDF file
        output_path (str): Path to output DOCX file
        maintain_formatting (bool): Whether to maintain original formatting
        extract_images (bool): Whether to extract images from PDF
        page_from (int): Start page (1-indexed), None for first page
        page_to (int): End page (1-indexed), None for last page

    Returns:
        bool: True if conversion successful, False otherwise
    """
    converter = None
    try:
        # Validate input file
        is_valid, error_msg = validate_pdf_file(input_path)
        if not is_valid:
            logger.error(f"Validation failed: {error_msg}")
            return False

        # Create output directory if it doesn't exist
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir, exist_ok=True)
                logger.info(f"Created output directory: {output_dir}")
            except Exception as e:
                logger.error(f"Failed to create output directory: {str(e)}")
                return False

        # Check if output path is writable
        if output_dir and not os.access(output_dir, os.W_OK):
            logger.error(f"Output directory is not writable: {output_dir}")
            return False

        logger.info(f"Starting conversion: '{input_path}' -> '{output_path}'")
        logger.info(f"Options: maintain_formatting={maintain_formatting}, extract_images={extract_images}")

        # Try pdf2docx first if available
        if HAS_PDF2DOCX:
            logger.info("Using pdf2docx library for conversion")
            try:
                converter = Converter(input_path)
                converter.convert(output_path)
                converter.close()
                converter = None

                # Verify output file was created
                if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                    output_size = os.path.getsize(output_path)
                    logger.info(f"Conversion completed successfully! Output size: {output_size / 1024:.2f}KB")
                    return True
                else:
                    logger.warning("pdf2docx conversion produced empty file, trying fallback")

            except Exception as e:
                logger.warning(f"pdf2docx conversion failed: {str(e)}, trying fallback method")

        # Use fallback method
        logger.info("Using fallback conversion method (PyMuPDF + python-docx)")
        success = convert_pdf_to_docx_fallback(input_path, output_path, page_from, page_to)

        if success:
            output_size = os.path.getsize(output_path)
            logger.info(f"Fallback conversion completed! Output size: {output_size / 1024:.2f}KB")

        return success

    except Exception as e:
        logger.error(f"Error during conversion: {str(e)}", exc_info=True)
        return False

    finally:
        # Ensure converter is properly closed
        if converter is not None:
            try:
                converter.close()
            except Exception as e:
                logger.warning(f"Error closing converter: {str(e)}")


def main():
    """Main entry point for the script"""
    if len(sys.argv) < 3:
        print("Usage: python3 pdf_to_docx.py <input_pdf> <output_docx> [options]", file=sys.stderr)
        print("\nOptions:", file=sys.stderr)
        print("  --no-maintain-formatting  Don't maintain original formatting", file=sys.stderr)
        print("  --no-extract-images       Don't extract images", file=sys.stderr)
        print("  --page-from N             Start page (1-indexed)", file=sys.stderr)
        print("  --page-to N               End page (1-indexed)", file=sys.stderr)
        print("\nExample:", file=sys.stderr)
        print("  python3 pdf_to_docx.py input.pdf output.docx", file=sys.stderr)
        print("  python3 pdf_to_docx.py input.pdf output.docx --page-from 1 --page-to 5", file=sys.stderr)
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    # Parse options
    maintain_formatting = True
    extract_images = True
    page_from = None
    page_to = None

    i = 3
    while i < len(sys.argv):
        arg = sys.argv[i]

        if arg == '--no-maintain-formatting':
            maintain_formatting = False
            logger.info("Formatting preservation disabled")
        elif arg == '--no-extract-images':
            extract_images = False
            logger.info("Image extraction disabled")
        elif arg == '--page-from':
            i += 1
            if i < len(sys.argv):
                try:
                    page_from = int(sys.argv[i])
                    logger.info(f"Start page set to: {page_from}")
                except ValueError:
                    logger.error(f"Invalid page number: {sys.argv[i]}")
                    sys.exit(1)
            else:
                logger.error("Error: --page-from requires a page number")
                sys.exit(1)
        elif arg == '--page-to':
            i += 1
            if i < len(sys.argv):
                try:
                    page_to = int(sys.argv[i])
                    logger.info(f"End page set to: {page_to}")
                except ValueError:
                    logger.error(f"Invalid page number: {sys.argv[i]}")
                    sys.exit(1)
            else:
                logger.error("Error: --page-to requires a page number")
                sys.exit(1)
        else:
            logger.warning(f"Unknown option: {arg}")

        i += 1

    logger.info("=" * 60)
    logger.info("PDF to DOCX Conversion Started")
    logger.info("=" * 60)

    # Perform conversion
    success = convert_pdf_to_docx(
        input_path,
        output_path,
        maintain_formatting=maintain_formatting,
        extract_images=extract_images,
        page_from=page_from,
        page_to=page_to
    )

    logger.info("=" * 60)
    if success:
        logger.info("Conversion completed successfully!")
    else:
        logger.error("Conversion failed!")
    logger.info("=" * 60)

    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
